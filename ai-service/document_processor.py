import os
from typing import List, Dict, Any
from datetime import datetime
import PyPDF2
from docx import Document

class DocumentProcessor:
    """
    Processes medical documents (PDF, DOCX) for RAG system
    
    Handles:
    - Text extraction from PDFs and Word documents
    - Chunking documents into manageable pieces
    - Metadata extraction
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Characters to overlap between chunks (for context)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from a document
        
        Args:
            file_path: Path to the document file
            filename: Original filename
            
        Returns:
            Document metadata
        """
        file_extension = os.path.splitext(filename)[1].lower()
        
        metadata = {
            "filename": filename,
            "upload_date": datetime.now().isoformat(),
            "file_type": file_extension
        }
        
        if file_extension == '.pdf':
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata["pages"] = len(pdf_reader.pages)
                # Try to get PDF metadata
                if pdf_reader.metadata:
                    metadata["title"] = pdf_reader.metadata.get('/Title', filename)
                    metadata["author"] = pdf_reader.metadata.get('/Author', 'Unknown')
        
        elif file_extension in ['.docx', '.doc']:
            doc = Document(file_path)
            metadata["pages"] = len(doc.paragraphs) // 20  # Rough estimate
        
        return metadata
    
    def extract_chunks(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract and chunk text from a document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of text chunks with metadata
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_chunks(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_docx_chunks(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_pdf_chunks(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text chunks from PDF with page tracking"""
        chunks = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        # Split page text into chunks
                        page_chunks = self._split_text_into_chunks(text)
                        for chunk_text in page_chunks:
                            chunks.append({
                                "text": chunk_text,
                                "page": page_num
                            })
                except Exception as e:
                    print(f"Error extracting page {page_num}: {str(e)}")
                    continue
        
        return chunks
    
    def _extract_docx_chunks(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text chunks from DOCX"""
        chunks = []
        
        doc = Document(file_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        
        if full_text.strip():
            chunk_texts = self._split_text_into_chunks(full_text)
            for i, chunk_text in enumerate(chunk_texts):
                chunks.append({
                    "text": chunk_text,
                    "page": f"section_{i+1}"
                })
        
        return chunks
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks
        
        This preserves context between chunks for better retrieval
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            # Get chunk of specified size
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence endings within the chunk
                for punct in ['. ', '.\n', '! ', '?\n']:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct != -1:
                        end = last_punct + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap if end < text_length else text_length
        
        return chunks
    
    def extract_medical_entities(self, text: str) -> Dict[str, List[str]]:
        """
        Extract medical entities from text (simple keyword-based approach)
        
        In production, use a medical NER model like:
        - scispacy (en_ner_bc5cdr_md)
        - Clinical BERT NER
        
        For now, this is a placeholder
        """
        # Common medical keywords to look for
        symptoms_keywords = ['pain', 'fever', 'cough', 'headache', 'nausea', 'fatigue']
        diagnosis_keywords = ['diagnosis', 'condition', 'disease', 'syndrome', 'disorder']
        treatment_keywords = ['treatment', 'therapy', 'medication', 'prescription', 'surgery']
        
        entities = {
            "symptoms": [],
            "diagnoses": [],
            "treatments": []
        }
        
        text_lower = text.lower()
        
        # Simple keyword matching (replace with NER in production)
        for keyword in symptoms_keywords:
            if keyword in text_lower:
                entities["symptoms"].append(keyword)
        
        for keyword in diagnosis_keywords:
            if keyword in text_lower:
                entities["diagnoses"].append(keyword)
        
        for keyword in treatment_keywords:
            if keyword in text_lower:
                entities["treatments"].append(keyword)
        
        return entities